from docx import Document

doc = Document('/home/ubuntu/upload/Examak_IEEE_Conference_Paper.docx')
print(f"Total paragraphs: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"P{i}: {p.text[:100]}...")
