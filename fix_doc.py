from docx import Document

doc = Document('/home/ubuntu/upload/Examak_IEEE_Conference_Paper.docx')

# Find the last non-empty paragraph
last_para_index = -1
for i in range(len(doc.paragraphs)-1, -1, -1):
    if doc.paragraphs[i].text.strip():
        last_para_index = i
        break

if last_para_index != -1:
    # Get the text of the last paragraph
    last_text = doc.paragraphs[last_para_index].text
    print(f"Removing last paragraph: {last_text}")
    
    # Append it to the previous non-empty paragraph to keep the content but remove the page break/orphan
    prev_para_index = -1
    for i in range(last_para_index - 1, -1, -1):
        if doc.paragraphs[i].text.strip():
            prev_para_index = i
            break
    
    if prev_para_index != -1:
        doc.paragraphs[prev_para_index].text += " " + last_text
        # Remove the last paragraph's element from the document
        p = doc.paragraphs[last_para_index]._element
        p.getparent().remove(p)
        print("Merged last paragraph into the previous one.")

doc.save('Examak_IEEE_Conference_Paper.docx')
